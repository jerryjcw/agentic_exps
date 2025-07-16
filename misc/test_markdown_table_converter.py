"""
Tests for the markdown table converter module.
"""

import pytest
import tempfile
import os
from docx import Document
from misc.markdown_table_converter import MarkdownTableConverter


class TestMarkdownTableConverter:
    """Test cases for MarkdownTableConverter class."""
    
    def setup_method(self):
        """Setup test fixtures."""
        self.converter = MarkdownTableConverter(max_column_width=40)
    
    def test_detect_simple_table(self):
        """Test detection of a simple markdown table."""
        text = """This is some text.

| Name | Age | City |
|------|-----|------|
| John | 25  | NYC  |
| Jane | 30  | LA   |

More text here."""
        
        tables = self.converter.detect_markdown_tables(text)
        assert len(tables) == 1
        
        start_line, table_lines = tables[0]
        assert start_line == 3  # Line where table starts
        assert len(table_lines) == 3  # Header + 2 data rows
        assert "| Name | Age | City |" in table_lines[0]
        assert "| John | 25  | NYC  |" in table_lines[1]
        assert "| Jane | 30  | LA   |" in table_lines[2]
    
    def test_detect_multiple_tables(self):
        """Test detection of multiple markdown tables."""
        text = """First table:

| A | B |
|---|---|
| 1 | 2 |

Some text between tables.

| X | Y | Z |
|---|---|---|
| a | b | c |
| d | e | f |

End of text."""
        
        tables = self.converter.detect_markdown_tables(text)
        assert len(tables) == 2
        
        # First table
        start_line1, table_lines1 = tables[0]
        assert start_line1 == 3
        assert len(table_lines1) == 2
        
        # Second table
        start_line2, table_lines2 = tables[1]
        assert start_line2 == 9
        assert len(table_lines2) == 3
    
    def test_detect_no_tables(self):
        """Test when no markdown tables are present."""
        text = """This is just regular text.
No tables here at all.
| This is not a complete table
Also not a table |"""
        
        tables = self.converter.detect_markdown_tables(text)
        assert len(tables) == 0
    
    def test_parse_table_row(self):
        """Test parsing of table rows."""
        # Simple row
        row = "| Name | Age | City |"
        cells = self.converter.parse_table_row(row)
        assert cells == ["Name", "Age", "City"]
        
        # Row with spaces
        row = "|  John  |  25  |  NYC  |"
        cells = self.converter.parse_table_row(row)
        assert cells == ["John", "25", "NYC"]
        
        # Row with empty cells
        row = "| Name |  | City |"
        cells = self.converter.parse_table_row(row)
        assert cells == ["Name", "", "City"]
    
    def test_wrap_text_short(self):
        """Test text wrapping for short text."""
        text = "Short text"
        wrapped = self.converter.wrap_text(text, 40)
        assert wrapped == "Short text"
    
    def test_wrap_text_long(self):
        """Test text wrapping for long text."""
        text = "This is a very long piece of text that should be wrapped"
        wrapped = self.converter.wrap_text(text, 20)
        lines = wrapped.split('\n')
        assert len(lines) > 1
        assert all(len(line) <= 20 for line in lines)
    
    def test_wrap_text_very_long_word(self):
        """Test text wrapping with very long words."""
        text = "Supercalifragilisticexpialidocious"
        wrapped = self.converter.wrap_text(text, 10)
        lines = wrapped.split('\n')
        assert len(lines) > 1
        assert all(len(line) <= 10 for line in lines)
    
    def test_convert_to_docx_with_tables(self):
        """Test conversion of markdown tables to DOCX."""
        # Create temporary input file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("""# Test Document

| Product | Price | Stock |
|---------|-------|-------|
| Apple   | $1.00 | 100   |
| Banana  | $0.50 | 200   |

Some text here.

| Name | Department | Salary |
|------|------------|--------|
| Alice | Engineering | $70000 |
| Bob   | Marketing   | $60000 |
""")
            input_file = f.name
        
        # Create temporary output file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_file = f.name
        
        try:
            # Convert
            success = self.converter.convert_to_docx(input_file, output_file)
            assert success
            
            # Verify output file exists
            assert os.path.exists(output_file)
            
            # Load and verify DOCX content
            doc = Document(output_file)
            
            # Check that document has content
            assert len(doc.paragraphs) > 0
            assert len(doc.tables) == 2  # Should have 2 tables
            
            # Check first table
            table1 = doc.tables[0]
            assert len(table1.rows) == 3  # Header + 2 data rows
            assert len(table1.columns) == 3  # 3 columns
            
            # Check second table
            table2 = doc.tables[1]
            assert len(table2.rows) == 3  # Header + 2 data rows
            assert len(table2.columns) == 3  # 3 columns
            
        finally:
            # Clean up
            os.unlink(input_file)
            os.unlink(output_file)
    
    def test_convert_to_docx_no_tables(self):
        """Test conversion when no tables are found."""
        # Create temporary input file without tables
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("""# Test Document

This is just regular text.
No tables here.
""")
            input_file = f.name
        
        # Create temporary output file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_file = f.name
        
        try:
            # Convert
            success = self.converter.convert_to_docx(input_file, output_file)
            assert not success  # Should return False when no tables found
            
        finally:
            # Clean up
            os.unlink(input_file)
            if os.path.exists(output_file):
                os.unlink(output_file)
    
    def test_convert_nonexistent_file(self):
        """Test conversion with non-existent input file."""
        success = self.converter.convert_to_docx("nonexistent.txt", "output.docx")
        assert not success
    
    def test_different_max_column_width(self):
        """Test converter with different max column widths."""
        converter_narrow = MarkdownTableConverter(max_column_width=10)
        converter_wide = MarkdownTableConverter(max_column_width=100)
        
        long_text = "This is a very long piece of text that needs wrapping"
        
        narrow_wrapped = converter_narrow.wrap_text(long_text, 10)
        wide_wrapped = converter_wide.wrap_text(long_text, 100)
        
        assert len(narrow_wrapped.split('\n')) > len(wide_wrapped.split('\n'))
    
    def test_table_with_empty_cells(self):
        """Test table detection with empty cells."""
        text = """| Name | Age | City |
|------|-----|------|
| John |     | NYC  |
|      | 30  | LA   |"""
        
        tables = self.converter.detect_markdown_tables(text)
        assert len(tables) == 1
        
        start_line, table_lines = tables[0]
        assert len(table_lines) == 3  # Header + 2 data rows
    
    def test_table_at_end_of_file(self):
        """Test table detection when table is at end of file."""
        text = """Some text.

| A | B |
|---|---|
| 1 | 2 |"""
        
        tables = self.converter.detect_markdown_tables(text)
        assert len(tables) == 1
        
        start_line, table_lines = tables[0]
        assert len(table_lines) == 2  # Header + 1 data row
    
    def test_malformed_table_lines(self):
        """Test handling of malformed table lines."""
        text = """| Name | Age |
|------|-----|
| John | 25  |
This is not a table line
| Jane | 30  |"""
        
        tables = self.converter.detect_markdown_tables(text)
        # The malformed line breaks the table into two separate tables
        assert len(tables) == 2
        
        start_line1, table_lines1 = tables[0]
        assert len(table_lines1) == 2  # Header + 1 data row
        assert "| Name | Age |" in table_lines1[0]
        assert "| John | 25  |" in table_lines1[1]
        
        start_line2, table_lines2 = tables[1]
        assert len(table_lines2) == 1  # Only 1 data row
        assert "| Jane | 30  |" in table_lines2[0]


if __name__ == '__main__':
    pytest.main([__file__])