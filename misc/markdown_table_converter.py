"""
Markdown Table to DOCX Converter

This module provides functionality to detect markdown tables in text files
and convert them to a visually formatted DOCX document.
"""

import re
from typing import List, Tuple, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.units import inch
import argparse
import sys
import os


class MarkdownTableConverter:
    """Converts markdown tables from text files to DOCX format."""
    
    def __init__(self, max_column_width: int = 40):
        """
        Initialize the converter.
        
        Args:
            max_column_width: Maximum characters per column (default: 40)
        """
        self.max_column_width = max_column_width
    
    def detect_markdown_tables(self, text: str) -> List[Tuple[int, List[str]]]:
        """
        Detect markdown tables in text using simple line-based detection.
        
        Args:
            text: Input text to scan for tables
            
        Returns:
            List of tuples containing (start_line_number, table_lines)
        """
        lines = text.split('\n')
        tables = []
        current_table = []
        table_start_line = None
        
        for i, line in enumerate(lines):
            stripped_line = line.strip()
            
            # Check if line starts with '|' (markdown table row)
            if stripped_line.startswith('|') and stripped_line.endswith('|'):
                if current_table == []:
                    table_start_line = i + 1  # 1-based line numbering
                current_table.append(stripped_line)
            else:
                # End of table or not a table line
                if current_table:
                    # Filter out separator lines (lines with only |, -, :, and spaces)
                    table_rows = []
                    for row in current_table:
                        if not re.match(r'^[\|\-\:\s]+$', row):
                            table_rows.append(row)
                    
                    if table_rows:  # Only add if there are actual data rows
                        tables.append((table_start_line, table_rows))
                    current_table = []
                    table_start_line = None
        
        # Handle table at end of file
        if current_table:
            table_rows = []
            for row in current_table:
                if not re.match(r'^[\|\-\:\s]+$', row):
                    table_rows.append(row)
            
            if table_rows:
                tables.append((table_start_line, table_rows))
        
        return tables
    
    def parse_table_row(self, row: str) -> List[str]:
        """
        Parse a markdown table row into individual cells.
        
        Args:
            row: Markdown table row string
            
        Returns:
            List of cell contents
        """
        # Remove leading and trailing |
        row = row.strip()
        if row.startswith('|'):
            row = row[1:]
        if row.endswith('|'):
            row = row[:-1]
        
        # Split by | and clean up
        cells = [cell.strip() for cell in row.split('|')]
        return cells
    
    def wrap_text(self, text: str, max_width: int) -> str:
        """
        Wrap text to fit within specified column width.
        
        Args:
            text: Text to wrap
            max_width: Maximum width per line
            
        Returns:
            Wrapped text with line breaks
        """
        if len(text) <= max_width:
            return text
        
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            if len(current_line + " " + word) <= max_width:
                if current_line:
                    current_line += " " + word
                else:
                    current_line = word
            else:
                if current_line:
                    lines.append(current_line)
                    current_line = word
                else:
                    # Word is longer than max_width, split it
                    while len(word) > max_width:
                        lines.append(word[:max_width])
                        word = word[max_width:]
                    current_line = word
        
        if current_line:
            lines.append(current_line)
        
        return '\n'.join(lines)
    
    def create_pdf_table(self, story: List, table_data: List[List[str]], 
                         table_number: int, start_line: int, styles) -> None:
        """
        Create a formatted table for PDF document.
        
        Args:
            story: List to append PDF elements to
            table_data: Parsed table data
            table_number: Table number for heading
            start_line: Line number where table starts
            styles: PDF styles
        """
        # Add heading
        heading = Paragraph(f'Table {table_number} (Line {start_line})', styles['Heading2'])
        story.append(heading)
        story.append(Spacer(1, 12))
        
        if not table_data:
            return
        
        # Wrap text in cells
        wrapped_data = []
        for row in table_data:
            wrapped_row = []
            for cell in row:
                wrapped_text = self.wrap_text(cell, self.max_column_width)
                wrapped_row.append(wrapped_text)
            wrapped_data.append(wrapped_row)
        
        # Create table
        table = Table(wrapped_data)
        
        # Style the table
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        
        story.append(table)
        story.append(Spacer(1, 12))
    
    def create_docx_table(self, doc: Document, table_data: List[List[str]], 
                         table_number: int, start_line: int) -> None:
        """
        Create a formatted table in the DOCX document.
        
        Args:
            doc: Document object
            table_data: Parsed table data
            table_number: Table number for heading
            start_line: Line number where table starts
        """
        # Add heading
        heading = doc.add_heading(f'Table {table_number} (Line {start_line})', level=2)
        
        # Determine number of columns
        max_cols = max(len(row) for row in table_data) if table_data else 0
        
        if max_cols == 0:
            return
        
        # Create table
        table = doc.add_table(rows=0, cols=max_cols)
        table.style = 'Light Grid Accent 1'
        
        # Add rows
        for row_data in table_data:
            row_cells = table.add_row().cells
            
            for i, cell_text in enumerate(row_data):
                if i < len(row_cells):
                    # Wrap text to fit column width
                    wrapped_text = self.wrap_text(cell_text, self.max_column_width)
                    row_cells[i].text = wrapped_text
                    
                    # Center align text
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing after table
        doc.add_paragraph()
    
    def convert_to_pdf(self, input_file: str, output_file: str) -> bool:
        """
        Convert markdown tables from input file to PDF format.
        
        Args:
            input_file: Path to input text file
            output_file: Path to output PDF file
            
        Returns:
            True if conversion successful, False otherwise
        """
        try:
            # Read input file
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Detect tables
            tables = self.detect_markdown_tables(text)
            
            if not tables:
                print(f"No markdown tables found in {input_file}")
                return False
            
            # Create PDF document
            doc = SimpleDocTemplate(output_file, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Add title
            title = Paragraph('Markdown Tables Export', styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # Add metadata
            info = Paragraph(f'Extracted from: {os.path.basename(input_file)}', styles['Normal'])
            story.append(info)
            tables_count = Paragraph(f'Total tables found: {len(tables)}', styles['Normal'])
            story.append(tables_count)
            story.append(Spacer(1, 20))
            
            # Process each table
            for table_num, (start_line, table_lines) in enumerate(tables, 1):
                table_data = []
                for line in table_lines:
                    cells = self.parse_table_row(line)
                    table_data.append(cells)
                
                self.create_pdf_table(story, table_data, table_num, start_line, styles)
            
            # Build PDF
            doc.build(story)
            print(f"Successfully converted {len(tables)} tables to {output_file}")
            return True
            
        except Exception as e:
            print(f"Error converting file: {e}")
            return False
    
    def convert_to_docx(self, input_file: str, output_file: str) -> bool:
        """
        Convert markdown tables from input file to DOCX format.
        
        Args:
            input_file: Path to input text file
            output_file: Path to output DOCX file
            
        Returns:
            True if conversion successful, False otherwise
        """
        try:
            # Read input file
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Detect tables
            tables = self.detect_markdown_tables(text)
            
            if not tables:
                print(f"No markdown tables found in {input_file}")
                return False
            
            # Create DOCX document
            doc = Document()
            doc.add_heading('Markdown Tables Export', level=1)
            doc.add_paragraph(f'Extracted from: {os.path.basename(input_file)}')
            doc.add_paragraph(f'Total tables found: {len(tables)}')
            doc.add_paragraph()
            
            # Process each table
            for table_num, (start_line, table_lines) in enumerate(tables, 1):
                table_data = []
                for line in table_lines:
                    cells = self.parse_table_row(line)
                    table_data.append(cells)
                
                self.create_docx_table(doc, table_data, table_num, start_line)
            
            # Save document
            doc.save(output_file)
            print(f"Successfully converted {len(tables)} tables to {output_file}")
            return True
            
        except Exception as e:
            print(f"Error converting file: {e}")
            return False


def main():
    """Command line interface for the markdown table converter."""
    parser = argparse.ArgumentParser(
        description='Convert markdown tables from text files to DOCX or PDF format'
    )
    parser.add_argument('input_file', help='Input text file containing markdown tables')
    parser.add_argument('output_file', help='Output file path (extension determines format: .docx or .pdf)')
    parser.add_argument('--max-width', type=int, default=40,
                       help='Maximum characters per column (default: 40)')
    
    args = parser.parse_args()
    
    # Check if input file exists
    if not os.path.exists(args.input_file):
        print(f"Error: Input file '{args.input_file}' not found")
        sys.exit(1)
    
    # Determine output format from file extension
    output_ext = os.path.splitext(args.output_file)[1].lower()
    
    # Create converter and run conversion
    converter = MarkdownTableConverter(max_column_width=args.max_width)
    
    if output_ext == '.pdf':
        success = converter.convert_to_pdf(args.input_file, args.output_file)
    elif output_ext == '.docx':
        success = converter.convert_to_docx(args.input_file, args.output_file)
    else:
        print(f"Error: Unsupported output format '{output_ext}'. Use .docx or .pdf")
        sys.exit(1)
    
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()